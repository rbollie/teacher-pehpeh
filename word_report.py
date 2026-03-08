"""
word_report.py — IBT Teacher Pehpeh Academic Word Report Generator
Generates a .docx with:
  - Cover / summary header
  - Charts: (1) Class avg by subject bar chart
            (2) Mom's Education × Kids in Home bar chart
            (3) HW & Quiz counts per semester grouped bar
  - Subject Breakdown table (class averages per subject)
  - Individual Averages table (per student × subject, sem 1/2)
  - Family Context table (Mom's Edu × Kids combos)
  - AI narrative
"""

import io
import datetime
import collections

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colours ────────────────────────────────────────────────────────────
IBT_GOLD      = RGBColor(0xD4, 0xA8, 0x43)
IBT_DARK_BLUE = RGBColor(0x0A, 0x0E, 0x1A)
IBT_BLUE      = RGBColor(0x1E, 0x3A, 0x6A)
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
RED           = RGBColor(0xC0, 0x39, 0x2B)    # Intervention
AMBER         = RGBColor(0xD4, 0x7E, 0x00)    # Monitor
GREEN         = RGBColor(0x1E, 0x8B, 0x4C)    # On Track
LIGHT_BLUE_BG = RGBColor(0xDA, 0xE8, 0xFC)
LIGHT_RED_BG  = RGBColor(0xF8, 0xD7, 0xDA)
LIGHT_AMB_BG  = RGBColor(0xFF, 0xF3, 0xCD)
LIGHT_GRN_BG  = RGBColor(0xD1, 0xEC, 0xE1)
HEADER_BG     = RGBColor(0x1E, 0x3A, 0x6A)

# Matplotlib hex colours
MC_BLUE    = "#1E3A6A"
MC_GOLD    = "#D4A843"
MC_RED     = "#C0392B"
MC_AMBER   = "#D47E00"
MC_GREEN   = "#1E8B4C"
MC_GRAY    = "#95A5A6"
MC_BG      = "#0A0E1A"
MC_TEXT    = "#DDDDDD"

SUBJ_PALETTE = [
    "#2E86C1","#D4A843","#1E8B4C","#C0392B","#8E44AD",
    "#D47E00","#17A589","#F39C12","#2C3E50","#E74C3C",
    "#27AE60","#3498DB","#E67E22","#9B59B6","#1ABC9C",
]

# ── Helpers ──────────────────────────────────────────────────────────────────
def _status(score):
    if score < 50:  return "Intervention"
    if score < 65:  return "Monitor"
    return "On Track"

def _status_rgb(score):
    if score < 50:  return RED,   LIGHT_RED_BG
    if score < 65:  return AMBER, LIGHT_AMB_BG
    return GREEN, LIGHT_GRN_BG

def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)

def _set_cell_borders(cell, color="1E3A6A"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def _hdr_para(cell, text, bold=True, font_size=10, color=WHITE):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color

def _data_para(cell, text, bold=False, font_size=9.5,
               color=IBT_DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color

def _add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = IBT_BLUE
    pPr = p._p.get_or_add_pPr()
    # bottom border = gold rule
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "D4A843")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p

def _img_from_fig(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight",
                facecolor=fig.get_facecolor(), dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf

# ── Chart generators ──────────────────────────────────────────────────────────

def _chart_subject_avg(df):
    """Bar chart: class average per subject."""
    subjs = sorted(df["subject"].unique())
    avgs  = [df[df["subject"]==s]["score"].mean() for s in subjs]
    colors = [MC_RED if a < 50 else (MC_AMBER if a < 65 else MC_GREEN) for a in avgs]

    fig, ax = plt.subplots(figsize=(8, 3.2), facecolor="white")
    bars = ax.bar(subjs, avgs, color=colors, edgecolor="white", linewidth=0.8, zorder=3)
    ax.set_ylim(0, 110)
    ax.set_ylabel("Class Average (%)", fontsize=9, color="#333")
    ax.set_title("Class Average by Subject", fontsize=11, fontweight="bold", color=MC_BLUE)
    ax.tick_params(axis="x", labelsize=8.5, rotation=20)
    ax.tick_params(axis="y", labelsize=8)
    ax.yaxis.grid(True, alpha=0.3, zorder=0)
    ax.set_axisbelow(True)
    ax.spines[["top","right"]].set_visible(False)
    for bar, val in zip(bars, avgs):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1.5,
                f"{val:.1f}", ha="center", va="bottom", fontsize=8, fontweight="bold",
                color="#333")
    # Legend
    patches = [
        mpatches.Patch(color=MC_GREEN, label="On Track (≥65)"),
        mpatches.Patch(color=MC_AMBER, label="Monitor (50–64)"),
        mpatches.Patch(color=MC_RED,   label="Intervention (<50)"),
    ]
    ax.legend(handles=patches, fontsize=7.5, loc="upper right",
              framealpha=0.8, edgecolor="#ccc")
    fig.tight_layout(pad=1.0)
    return _img_from_fig(fig)


def _chart_family_context(df, students):
    """Bar chart: avg score by Mom's Education × Kids in Home combo."""
    if not students:
        return None
    hs_map  = {"No HS": "No High School", "HS Grad": "High School Grad", "Unknown": "Unknown"}
    sib_map = {"0-4": "0–4 kids", "5-8": "5–8 kids", "8+": "8+ kids", "Unknown": "Unknown"}
    fam_map = {}
    stu_lookup = {s["name"]: s for s in students}
    for stn in df["student"].unique():
        fp  = stu_lookup.get(stn, {})
        mom = (fp.get("mom") or "Unknown").strip() or "Unknown"
        sib = (fp.get("sib") or "Unknown").strip() or "Unknown"
        avg = df[df["student"] == stn]["score"].mean()
        fam_map.setdefault((mom, sib), []).append(avg)
    if not fam_map:
        return None

    mom_ord = ["No HS","HS Grad","Unknown"]
    sib_ord = ["0-4","5-8","8+","Unknown"]
    rows = sorted(fam_map.items(),
        key=lambda kv: (mom_ord.index(kv[0][0]) if kv[0][0] in mom_ord else 99,
                        sib_ord.index(kv[0][1]) if kv[0][1] in sib_ord else 99))
    labels = [f"{hs_map.get(m,m)}\n{sib_map.get(k,k)}" for (m,k),_ in rows]
    avgs   = [round(sum(v)/len(v), 1) for _,v in rows]
    colors = [MC_RED if a < 50 else (MC_AMBER if a < 65 else MC_GREEN) for a in avgs]

    fig, ax = plt.subplots(figsize=(8, 3.4), facecolor="white")
    bars = ax.bar(range(len(labels)), avgs, color=colors, edgecolor="white", linewidth=0.8, zorder=3)
    ax.set_xticks(range(len(labels)))
    ax.set_xticklabels(labels, fontsize=8, rotation=0, ha="center")
    ax.set_ylim(0, 110)
    ax.set_ylabel("Average Score (%)", fontsize=9, color="#333")
    ax.set_title("Average Score by Mom's Education & Kids in Home", fontsize=11,
                 fontweight="bold", color=MC_BLUE)
    ax.yaxis.grid(True, alpha=0.3, zorder=0)
    ax.set_axisbelow(True)
    ax.spines[["top","right"]].set_visible(False)
    for bar, val in zip(bars, avgs):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1.5,
                f"{val:.1f}", ha="center", va="bottom", fontsize=8, fontweight="bold",
                color="#333")
    patches = [
        mpatches.Patch(color=MC_GREEN, label="On Track (≥65)"),
        mpatches.Patch(color=MC_AMBER, label="Monitor (50–64)"),
        mpatches.Patch(color=MC_RED,   label="Intervention (<50)"),
    ]
    ax.legend(handles=patches, fontsize=7.5, loc="upper right", framealpha=0.8, edgecolor="#ccc")
    fig.tight_layout(pad=1.0)
    return _img_from_fig(fig)


def _chart_hw_quiz(subject_data):
    """Grouped bar: HW & Quiz count per subject per semester."""
    if not subject_data:
        return None
    subjs = sorted(subject_data.keys())
    sems  = ["Semester 1", "Semester 2"]
    key_map = {
        ("Semester 1","Homework"): "hw1_count",
        ("Semester 1","Quiz"):     "qz1_count",
        ("Semester 2","Homework"): "hw2_count",
        ("Semester 2","Quiz"):     "qz2_count",
    }
    # Aggregate across all students
    agg = {(sem, atype): [] for sem in sems for atype in ["Homework","Quiz"]}
    for subj in subjs:
        sdata = subject_data.get(subj, {})
        for sem in sems:
            for atype in ["Homework","Quiz"]:
                counts = [v.get(key_map[(sem,atype)], 0) or 0
                          for v in sdata.values()]
                agg[(sem,atype)].append(int(np.median(counts)) if counts else 0)

    x = np.arange(len(subjs))
    w = 0.2
    offsets = [-1.5*w, -0.5*w, 0.5*w, 1.5*w]
    bar_defs = [
        ("Semester 1","Homework", "#2E86C1"),
        ("Semester 1","Quiz",     "#D4A843"),
        ("Semester 2","Homework", "#1E8B4C"),
        ("Semester 2","Quiz",     "#C0392B"),
    ]
    fig, ax = plt.subplots(figsize=(8, 3.4), facecolor="white")
    for (sem, atype, color), off in zip(bar_defs, offsets):
        vals = agg[(sem, atype)]
        ax.bar(x + off, vals, width=w, label=f"{sem} – {atype}",
               color=color, edgecolor="white", linewidth=0.6, zorder=3)
    ax.set_xticks(x)
    ax.set_xticklabels(subjs, fontsize=8.5, rotation=20, ha="right")
    ax.set_ylabel("Median Count", fontsize=9, color="#333")
    ax.set_title("Homework & Quiz Counts per Subject by Semester", fontsize=11,
                 fontweight="bold", color=MC_BLUE)
    ax.yaxis.grid(True, alpha=0.3, zorder=0)
    ax.set_axisbelow(True)
    ax.spines[["top","right"]].set_visible(False)
    ax.legend(fontsize=7.5, ncol=2, loc="upper right", framealpha=0.8, edgecolor="#ccc")
    fig.tight_layout(pad=1.0)
    return _img_from_fig(fig)


# ── Table builders ────────────────────────────────────────────────────────────

def _table_subject_breakdown(doc, df, subject_data):
    """Subject Breakdown — class averages per subject."""
    _add_section_heading(doc, "📚 Subject Breakdown — Class Averages", level=2)

    subjs = sorted(df["subject"].unique())
    has_s1 = False; has_s2 = False
    rows_data = []
    for subj in subjs:
        sd2 = df[df["subject"] == subj]
        cls_avg = sd2["score"].mean()
        sdata   = (subject_data or {}).get(subj, {})
        s1_vals = [v["s1_avg"] for v in sdata.values() if v.get("s1_avg") is not None]
        s2_vals = [v["s2_avg"] for v in sdata.values() if v.get("s2_avg") is not None]
        s1_avg  = round(sum(s1_vals)/len(s1_vals), 1) if s1_vals else None
        s2_avg  = round(sum(s2_vals)/len(s2_vals), 1) if s2_vals else None
        if s1_avg is not None: has_s1 = True
        if s2_avg is not None: has_s2 = True
        rows_data.append((subj, round(cls_avg,1), s1_avg, s2_avg,
                          sd2["student"].nunique(), round(sd2["score"].min(),1)))

    # Build columns list
    col_hdrs = ["Subject", "Class Avg"]
    if has_s1: col_hdrs.append("Sem 1 Avg")
    if has_s2: col_hdrs.append("Sem 2 Avg")
    col_hdrs += ["# Students", "Lowest", "Status"]

    ncols = len(col_hdrs)
    col_widths_in = {
        "Subject": 1.7, "Class Avg": 1.0, "Sem 1 Avg": 1.0,
        "Sem 2 Avg": 1.0, "# Students": 0.85, "Lowest": 0.9, "Status": 1.15,
    }
    widths = [col_widths_in[h] for h in col_hdrs]

    tbl = doc.add_table(rows=1 + len(rows_data), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header row
    hdr = tbl.rows[0]
    for ci, (hdr_text, w) in enumerate(zip(col_hdrs, widths)):
        cell = hdr.cells[ci]
        cell.width = Inches(w)
        _set_cell_bg(cell, HEADER_BG)
        _set_cell_borders(cell, "2E5A9C")
        _hdr_para(cell, hdr_text)

    # Data rows
    for ri, (subj, cls_avg, s1, s2, n_stu, lowest) in enumerate(rows_data):
        row = tbl.rows[ri + 1]
        fg, bg = _status_rgb(cls_avg)
        row_vals = [subj, f"{cls_avg}/100"]
        if has_s1: row_vals.append(f"{s1}/100" if s1 else "—")
        if has_s2: row_vals.append(f"{s2}/100" if s2 else "—")
        row_vals += [str(n_stu), f"{lowest}/100", _status(cls_avg)]
        row_bg = LIGHT_BLUE_BG if ri % 2 == 0 else RGBColor(0xF5, 0xF8, 0xFF)

        for ci, (val, w) in enumerate(zip(row_vals, widths)):
            cell = row.cells[ci]
            cell.width = Inches(w)
            is_score = col_hdrs[ci] in ("Class Avg","Sem 1 Avg","Sem 2 Avg","Lowest","Status")
            _set_cell_bg(cell, bg if is_score else row_bg)
            _set_cell_borders(cell, "ADC6E5")
            _data_para(cell, val, bold=(col_hdrs[ci]=="Class Avg"),
                       color=fg if is_score else IBT_DARK_BLUE,
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def _table_individual_averages(doc, df, subject_data):
    """Individual averages per student × subject, grouped by student with semester columns."""
    _add_section_heading(doc, "👤 Individual Averages by Subject & Semester", level=2)

    sdata_all = subject_data or {}
    has_sem = any(
        v.get("s1_avg") is not None or v.get("s2_avg") is not None
        for sd in sdata_all.values() for v in sd.values()
    )

    if has_sem:
        col_hdrs = ["Student", "Subject", "Sem 1", "Sem 2", "Overall", "Status"]
        widths    = [1.4, 1.6, 0.95, 0.95, 1.0, 1.1]
        all_stus  = sorted({stu for sd in sdata_all.values() for stu in sd})
        all_subjs = sorted(sdata_all.keys())

        # Pre-build data rows
        ind_rows = []
        for stu in all_stus:
            for subj in all_subjs:
                sv   = sdata_all.get(subj, {}).get(stu, {})
                s1   = sv.get("s1_avg"); s2 = sv.get("s2_avg")
                if s1 is None and s2 is None: continue
                vals = [v for v in [s1, s2] if v is not None]
                ov   = round(sum(vals)/len(vals), 1)
                ind_rows.append({
                    "student": stu, "subject": subj,
                    "s1": f"{s1:.1f}/100" if s1 is not None else "—",
                    "s2": f"{s2:.1f}/100" if s2 is not None else "—",
                    "overall": f"{ov:.1f}/100", "ov_val": ov,
                })
    else:
        col_hdrs = ["Student", "Subject", "Avg Score", "Status"]
        widths    = [1.5, 1.8, 1.2, 1.2]
        ind_rows  = []
        for stu in sorted(df["student"].unique()):
            for subj in sorted(df["subject"].unique()):
                sdf = df[(df["student"]==stu) & (df["subject"]==subj)]
                if sdf.empty: continue
                avg = round(sdf["score"].mean(), 1)
                ind_rows.append({
                    "student": stu, "subject": subj,
                    "avg": f"{avg}/100", "ov_val": avg,
                })

    if not ind_rows:
        doc.add_paragraph("No individual data available.")
        return

    ncols = len(col_hdrs)
    total_rows = len(ind_rows) + 1  # +1 header
    tbl = doc.add_table(rows=total_rows, cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    # Header row
    for ci, (h, w) in enumerate(zip(col_hdrs, widths)):
        cell = tbl.rows[0].cells[ci]
        cell.width = Inches(w)
        _set_cell_bg(cell, HEADER_BG)
        _set_cell_borders(cell, "2E5A9C")
        _hdr_para(cell, h)

    # Data rows — track student boundaries for shading groups
    stu_color_idx = {}
    stu_counter = 0
    for ri, row_d in enumerate(ind_rows):
        stu = row_d["student"]
        if stu not in stu_color_idx:
            stu_color_idx[stu] = stu_counter
            stu_counter += 1
        row_bg = LIGHT_BLUE_BG if stu_color_idx[stu] % 2 == 0 else RGBColor(0xF5, 0xF8, 0xFF)
        fg, score_bg = _status_rgb(row_d["ov_val"])

        tbl_row = tbl.rows[ri + 1]

        if has_sem:
            vals = [stu, row_d["subject"], row_d["s1"], row_d["s2"],
                    row_d["overall"], _status(row_d["ov_val"])]
        else:
            vals = [stu, row_d["subject"], row_d["avg"], _status(row_d["ov_val"])]

        for ci, (val, w) in enumerate(zip(vals, widths)):
            cell = tbl_row.cells[ci]
            cell.width = Inches(w)
            is_score_col = col_hdrs[ci] in ("Sem 1","Sem 2","Overall","Avg Score","Status")
            is_stu_col   = col_hdrs[ci] == "Student"
            _set_cell_bg(cell, score_bg if is_score_col else row_bg)
            _set_cell_borders(cell, "ADC6E5")
            _data_para(cell, val,
                       bold=is_stu_col,
                       color=IBT_BLUE if is_stu_col else (fg if is_score_col else IBT_DARK_BLUE),
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci <= 1 else WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()


def _table_family_context(doc, df, students):
    """Family Context: Mom's Education × Kids combos → avg score."""
    if not students:
        return
    _add_section_heading(doc, "👨‍👩‍👧 Family Context — Performance by Background", level=2)

    hs_map  = {"No HS": "No High School", "HS Grad": "High School Grad", "Unknown": "Unknown"}
    sib_map = {"0-4": "0–4 kids", "5-8": "5–8 kids", "8+": "8+ kids", "Unknown": "Unknown"}
    fam_map = {}
    stu_lk  = {s["name"]: s for s in students}
    for stn in df["student"].unique():
        fp  = stu_lk.get(stn, {})
        mom = (fp.get("mom") or "Unknown").strip() or "Unknown"
        sib = (fp.get("sib") or "Unknown").strip() or "Unknown"
        avg = df[df["student"]==stn]["score"].mean()
        fam_map.setdefault((mom, sib), []).append(avg)

    if not fam_map:
        return

    mom_ord = ["No HS","HS Grad","Unknown"]
    sib_ord = ["0-4","5-8","8+","Unknown"]
    rows = sorted(fam_map.items(),
        key=lambda kv: (mom_ord.index(kv[0][0]) if kv[0][0] in mom_ord else 99,
                        sib_ord.index(kv[0][1]) if kv[0][1] in sib_ord else 99))

    col_hdrs = ["Mom's Education", "Kids in Home", "# Students", "Avg Score", "Status"]
    widths   = [1.8, 1.3, 1.0, 1.0, 1.2]
    ncols    = len(col_hdrs)

    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    for ci, (h, w) in enumerate(zip(col_hdrs, widths)):
        cell = tbl.rows[0].cells[ci]
        cell.width = Inches(w)
        _set_cell_bg(cell, HEADER_BG)
        _set_cell_borders(cell, "2E5A9C")
        _hdr_para(cell, h)

    for ri, ((mom, sib), vals) in enumerate(rows):
        avg = round(sum(vals)/len(vals), 1)
        fg, bg = _status_rgb(avg)
        row_bg = LIGHT_BLUE_BG if ri % 2 == 0 else RGBColor(0xF5, 0xF8, 0xFF)
        row_vals = [hs_map.get(mom, mom), sib_map.get(sib, sib),
                    len(vals), f"{avg}/100", _status(avg)]
        tbl_row = tbl.rows[ri + 1]
        for ci, (val, w) in enumerate(zip(row_vals, widths)):
            cell = tbl_row.cells[ci]
            cell.width = Inches(w)
            is_score = col_hdrs[ci] in ("Avg Score","Status")
            _set_cell_bg(cell, bg if is_score else row_bg)
            _set_cell_borders(cell, "ADC6E5")
            _data_para(cell, val,
                       bold=(col_hdrs[ci]=="Avg Score"),
                       color=fg if is_score else IBT_DARK_BLUE,
                       align=WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


# ── Main entry point ──────────────────────────────────────────────────────────

def generate_academic_word_report(
    grade_history,
    students,
    school_label,
    grade_level,
    ai_narrative="",
    subject_data=None,
):
    """
    Returns bytes of a .docx report including charts and tables.

    Parameters
    ----------
    grade_history : list of dicts  (student, subject, score, date, ...)
    students      : list of dicts  (name, mom, sib, ...)
    school_label  : str
    grade_level   : str
    ai_narrative  : str
    subject_data  : dict  subject → student → {s1_avg, s2_avg, hw1_count, ...}
    """
    if not grade_history:
        raise ValueError("No grade history to report on.")

    df = pd.DataFrame(grade_history)
    df["score"] = pd.to_numeric(df["score"], errors="coerce")
    df = df.dropna(subset=["score"])
    if df.empty:
        raise ValueError("No valid scores in grade history.")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width   = Cm(21.59)   # 8.5 in
        section.page_height  = Cm(27.94)   # 11 in
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # ── Cover block ───────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("IBT ACADEMIC PERFORMANCE REPORT")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = IBT_BLUE

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run(
        f"Institute of Basic Technology  •  {school_label or 'Class Report'}")
    run2.font.size = Pt(11)
    run2.font.color.rgb = IBT_DARK_BLUE

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now_str = datetime.datetime.now().strftime("%B %d, %Y")
    run3 = meta_p.add_run(
        f"Grade Level: {grade_level or 'All'}  •  Generated: {now_str}")
    run3.font.size = Pt(9.5)
    run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    meta_p.paragraph_format.space_after = Pt(10)

    # ── Summary KPI row ────────────────────────────────────────────────────────
    avg_all    = df["score"].mean()
    n_stu      = df["student"].nunique()
    n_subj     = df["subject"].nunique()
    below50    = df.groupby("student")["score"].mean().lt(50).sum()
    below65    = df.groupby("student")["score"].mean().between(50, 64.99).sum()
    on_track   = n_stu - below50 - below65

    kpi_tbl = doc.add_table(rows=1, cols=4)
    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpi_tbl.style = "Table Grid"
    kpi_data = [
        ("Class Average", f"{avg_all:.1f}%", avg_all),
        ("Students", str(n_stu), None),
        ("Subjects", str(n_subj), None),
        ("Need Intervention", str(below50), 0 if below50 == 0 else 30),
    ]
    for ci, (label, value, score_hint) in enumerate(kpi_data):
        cell = kpi_tbl.rows[0].cells[ci]
        cell.width = Inches(1.7)
        if score_hint is not None:
            bg = LIGHT_GRN_BG if score_hint >= 65 else (LIGHT_AMB_BG if score_hint >= 50 else LIGHT_RED_BG)
        else:
            bg = LIGHT_BLUE_BG
        _set_cell_bg(cell, bg)
        _set_cell_borders(cell, "ADC6E5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}\n").bold = False
        p.runs[-1].font.size = Pt(8.5)
        p.runs[-1].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        vrun = p.add_run(value)
        vrun.bold = True
        vrun.font.size = Pt(16)
        vrun.font.color.rgb = IBT_BLUE
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Section 1: Charts ─────────────────────────────────────────────────────
    _add_section_heading(doc, "📊 Performance Charts", level=1)

    # Chart 1: Subject averages
    try:
        img1 = _chart_subject_avg(df)
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img1 = p_img1.add_run()
        run_img1.add_picture(img1, width=Inches(6.0))
        cap1 = doc.add_paragraph("Figure 1 — Class average score per subject. "
                                  "Green = On Track (≥65), Amber = Monitor (50–64), Red = Intervention (<50).")
        cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap1.runs[0].font.size = Pt(8); cap1.runs[0].font.italic = True
        cap1.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()
    except Exception:
        pass

    # Chart 2: Family context
    try:
        img2 = _chart_family_context(df, students)
        if img2:
            p_img2 = doc.add_paragraph()
            p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img2 = p_img2.add_run()
            run_img2.add_picture(img2, width=Inches(6.0))
            cap2 = doc.add_paragraph(
                "Figure 2 — Average score grouped by mother's education level and number of siblings at home.")
            cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap2.runs[0].font.size = Pt(8); cap2.runs[0].font.italic = True
            cap2.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()
    except Exception:
        pass

    # Chart 3: HW/Quiz counts
    try:
        img3 = _chart_hw_quiz(subject_data)
        if img3:
            p_img3 = doc.add_paragraph()
            p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img3 = p_img3.add_run()
            run_img3.add_picture(img3, width=Inches(6.0))
            cap3 = doc.add_paragraph(
                "Figure 3 — Median homework and quiz counts per subject, split by semester.")
            cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap3.runs[0].font.size = Pt(8); cap3.runs[0].font.italic = True
            cap3.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            doc.add_paragraph()
    except Exception:
        pass

    # ── Section 2: Subject breakdown table ───────────────────────────────────
    _add_section_heading(doc, "📋 Data Tables", level=1)
    try:
        _table_subject_breakdown(doc, df, subject_data)
    except Exception:
        pass

    # ── Section 3: Individual averages table ──────────────────────────────────
    try:
        _table_individual_averages(doc, df, subject_data)
    except Exception:
        pass

    # ── Section 4: Family context table ──────────────────────────────────────
    try:
        _table_family_context(doc, df, students)
    except Exception:
        pass

    # ── Section 5: AI Narrative ───────────────────────────────────────────────
    if ai_narrative and ai_narrative.strip():
        _add_section_heading(doc, "🤖 AI Analysis Narrative", level=1)
        # Split narrative by line, add paragraphs
        for line in ai_narrative.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            # Simple bold detection for markdown-style headers (## or **)
            if line.startswith("##"):
                p = doc.add_paragraph()
                run = p.add_run(line.lstrip("#").strip())
                run.bold = True; run.font.size = Pt(11)
                run.font.color.rgb = IBT_BLUE
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True; run.font.size = Pt(10)
                run.font.color.rgb = IBT_DARK_BLUE
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(4)
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = IBT_DARK_BLUE

    # ── Footer note ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        f"Generated by Teacher Pehpeh — IBT (Institute of Basic Technology)  •  "
        f"Confidential — {now_str}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Serialise ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
